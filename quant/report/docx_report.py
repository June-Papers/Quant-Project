from pathlib import Path
from typing import Dict, Optional
import pandas as pd

try:
    import docx
    from docx.shared import Inches
except ImportError:
    docx = None


def dataframe_to_markdown(df: pd.DataFrame) -> str:
    return df.to_markdown(index=True)


class ReportAssembler:
    def __init__(self, output_dir: Path):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def write_markdown(
        self,
        path: Path,
        title: str,
        overview: str,
        tables: Dict[str, pd.DataFrame],
        images: Dict[str, Path],
    ) -> Path:
        content = [f"# {title}", ""]
        content.append(overview)
        content.append("")

        for section, df in tables.items():
            content.append(f"## {section}")
            content.append("")
            content.append(dataframe_to_markdown(df))
            content.append("")

        for label, image_path in images.items():
            content.append(f"### {label.replace('_', ' ').title()}")
            content.append(f"![{label}]({image_path.name})")
            content.append("")

        out_path = self.output_dir / path
        out_path.write_text("\n".join(content), encoding="utf-8")
        return out_path

    def write_docx(
        self,
        path: Path,
        title: str,
        overview: str,
        tables: Dict[str, pd.DataFrame],
        images: Dict[str, Path],
    ) -> Optional[Path]:
        if docx is None:
            return None

        doc = docx.Document()
        doc.add_heading(title, level=1)
        doc.add_paragraph(overview)

        for section, df in tables.items():
            doc.add_heading(section, level=2)
            table = doc.add_table(rows=1, cols=len(df.columns) + 1)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = df.index.name or ""
            for idx, col in enumerate(df.columns, start=1):
                hdr_cells[idx].text = str(col)

            for idx, (index_label, row) in enumerate(df.iterrows(), start=1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(index_label)
                for col_idx, value in enumerate(row, start=1):
                    row_cells[col_idx].text = str(round(value, 6) if isinstance(value, (int, float)) else value)

        for label, image_path in images.items():
            doc.add_heading(label.replace('_', ' ').title(), level=2)
            doc.add_picture(str(image_path), width=Inches(6))

        out_path = self.output_dir / path
        doc.save(out_path)
        return out_path
